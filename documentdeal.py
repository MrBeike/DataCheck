#!/usr/bin/python
# -*-coding:utf-8-*-

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32ui
import re

# # 1表示打开文件对话框
# dlg = win32ui.CreateFileDialog(1)
# # 设置打开文件对话框中的初始显示目录
# dlg.SetOFNInitialDir('C:/')
# dlg.DoModal()
# # 获取选择的文件名称
# filename = dlg.GetPathName()
# print(filename)
# # 打开指定文件
# document = Document(filename)
document = Document(r"c:\Users\XCB\Desktop\池州中支周报.doc")
paragraphs = document.paragraphs
lines = len(paragraphs)
for i in range(5, lines):
    contents = paragraphs[i].text
    # print(contents)
    if contents == re.compile(r"/\【(.+)/\】$"):
        # print(contents)
        continue
    else:
        print(contents)
        print(paragraphs[i].runs)
